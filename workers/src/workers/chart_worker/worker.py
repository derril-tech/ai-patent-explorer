import asyncio
import json
import logging
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas

from ..base import BaseWorker
from ...utils.database import DatabaseClient
from ...utils.storage import StorageClient

logger = logging.getLogger(__name__)


class ChartWorker(BaseWorker):
    """Worker for building claim charts and exports."""
    
    def __init__(self, nats_url: str = "nats://localhost:4222"):
        super().__init__(nats_url)
        self.db = DatabaseClient()
        self.storage = StorageClient()
        
        logger.info("ChartWorker initialized")
    
    async def start(self):
        """Start the chart worker."""
        await super().start()
        await self.db.connect()
        await self.storage.connect()
        
        # Subscribe to chart generation requests
        await self.subscribe("chart.generate", self.handle_chart_request)
        await self.subscribe("export.bundle", self.handle_export_request)
        
        logger.info("ChartWorker started and listening for requests")
    
    async def stop(self):
        """Stop the chart worker."""
        await self.db.disconnect()
        await self.storage.disconnect()
        await super().stop()
    
    async def handle_chart_request(self, msg):
        """Handle claim chart generation requests."""
        try:
            data = json.loads(msg.data.decode())
            chart_id = data.get('chart_id')
            patent_id = data.get('patent_id')
            claim_num = data.get('claim_num')
            chart_type = data.get('chart_type', 'docx')  # docx or pdf
            include_alignments = data.get('include_alignments', True)
            include_novelty = data.get('include_novelty', True)
            
            if not all([chart_id, patent_id, claim_num]):
                logger.error("Missing required fields in chart request")
                return
            
            logger.info(f"Processing chart request {chart_id} for patent {patent_id}, claim {claim_num}")
            
            # Generate claim chart
            chart_data = await self.generate_claim_chart(
                patent_id, claim_num, include_alignments, include_novelty
            )
            
            # Create document
            if chart_type == 'docx':
                file_path = await self.create_docx_chart(chart_id, chart_data)
            else:
                file_path = await self.create_pdf_chart(chart_id, chart_data)
            
            # Upload to storage
            s3_url = await self.upload_chart_to_storage(file_path, chart_id, chart_type)
            
            # Clean up local file
            os.remove(file_path)
            
            # Publish completion event
            await self.publish("chart.complete", {
                "chart_id": chart_id,
                "patent_id": patent_id,
                "claim_num": claim_num,
                "chart_type": chart_type,
                "file_url": s3_url,
                "status": "success"
            })
            
        except Exception as e:
            logger.error(f"Error processing chart request: {e}")
            await self.publish("chart.error", {
                "chart_id": data.get('chart_id'),
                "error": str(e)
            })
    
    async def handle_export_request(self, msg):
        """Handle export bundle requests."""
        try:
            data = json.loads(msg.data.decode())
            export_id = data.get('export_id')
            patent_ids = data.get('patent_ids', [])
            export_type = data.get('export_type', 'zip')  # zip or pdf
            include_charts = data.get('include_charts', True)
            include_alignments = data.get('include_alignments', True)
            include_novelty = data.get('include_novelty', True)
            
            if not all([export_id, patent_ids]):
                logger.error("Missing required fields in export request")
                return
            
            logger.info(f"Processing export request {export_id} for {len(patent_ids)} patents")
            
            # Generate export bundle
            bundle_path = await self.create_export_bundle(
                export_id, patent_ids, export_type, include_charts, 
                include_alignments, include_novelty
            )
            
            # Upload to storage
            s3_url = await self.upload_export_to_storage(bundle_path, export_id, export_type)
            
            # Clean up local file
            os.remove(bundle_path)
            
            # Publish completion event
            await self.publish("export.complete", {
                "export_id": export_id,
                "patent_ids": patent_ids,
                "export_type": export_type,
                "file_url": s3_url,
                "status": "success"
            })
            
        except Exception as e:
            logger.error(f"Error processing export request: {e}")
            await self.publish("export.error", {
                "export_id": data.get('export_id'),
                "error": str(e)
            })
    
    async def generate_claim_chart(
        self, 
        patent_id: str, 
        claim_num: int, 
        include_alignments: bool = True,
        include_novelty: bool = True
    ) -> Dict[str, Any]:
        """Generate claim chart data."""
        try:
            # Get patent and claim data
            patent = await self.db.get_patent(patent_id)
            claim = await self.db.get_claim(patent_id, claim_num)
            
            if not patent or not claim:
                raise ValueError(f"Patent or claim not found: {patent_id}, {claim_num}")
            
            chart_data = {
                'patent': patent,
                'claim': claim,
                'alignments': [],
                'novelty': None,
                'generated_at': datetime.utcnow().isoformat()
            }
            
            # Get alignments if requested
            if include_alignments:
                alignments = await self.db.get_claim_alignments(patent_id, claim_num)
                chart_data['alignments'] = alignments
            
            # Get novelty data if requested
            if include_novelty:
                novelty = await self.db.get_novelty_score(patent_id, claim_num)
                chart_data['novelty'] = novelty
            
            return chart_data
            
        except Exception as e:
            logger.error(f"Error generating claim chart: {e}")
            raise
    
    async def create_docx_chart(self, chart_id: str, chart_data: Dict[str, Any]) -> str:
        """Create a DOCX claim chart."""
        try:
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Claim Chart', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add patent information
            patent = chart_data['patent']
            claim = chart_data['claim']
            
            doc.add_heading('Patent Information', level=1)
            patent_info = doc.add_paragraph()
            patent_info.add_run(f"Title: {patent['title']}\n")
            patent_info.add_run(f"Publication Number: {patent['pub_number']}\n")
            patent_info.add_run(f"Priority Date: {patent['prio_date']}\n")
            patent_info.add_run(f"Assignee(s): {', '.join(patent['assignees'])}\n")
            patent_info.add_run(f"Inventors: {', '.join(patent['inventors'])}\n")
            
            # Add claim text
            doc.add_heading(f'Claim {claim["claim_number"]}', level=1)
            claim_text = doc.add_paragraph(claim['text'])
            claim_text.style = 'Quote'
            
            # Add novelty information if available
            if chart_data.get('novelty'):
                novelty = chart_data['novelty']
                doc.add_heading('Novelty Analysis', level=1)
                
                novelty_info = doc.add_paragraph()
                novelty_info.add_run(f"Novelty Score: {novelty['novelty_score']:.2f}\n")
                novelty_info.add_run(f"Obviousness Score: {novelty['obviousness_score']:.2f}\n")
                novelty_info.add_run(f"Confidence Band: {novelty['confidence_band']}\n")
                
                # Add clause-level details
                if novelty.get('clause_details'):
                    doc.add_heading('Clause-Level Analysis', level=2)
                    for clause in novelty['clause_details']:
                        clause_para = doc.add_paragraph()
                        clause_para.add_run(f"Clause {clause['clause_index']}: {clause['clause_text']}\n")
                        clause_para.add_run(f"Novelty Score: {clause['novelty_score']:.2f} (Confidence: {clause['confidence']})\n")
            
            # Add alignments if available
            if chart_data.get('alignments'):
                doc.add_heading('Reference Alignments', level=1)
                
                # Create alignment table
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Add headers
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Reference Patent'
                header_cells[1].text = 'Reference Clause'
                header_cells[2].text = 'Similarity Score'
                header_cells[3].text = 'Alignment Type'
                header_cells[4].text = 'Overlap Details'
                
                # Add alignment data
                for alignment in chart_data['alignments']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = alignment.get('reference_patent_title', 'Unknown')
                    row_cells[1].text = alignment['reference_clause_text'][:100] + '...' if len(alignment['reference_clause_text']) > 100 else alignment['reference_clause_text']
                    row_cells[2].text = f"{alignment['similarity_score']:.3f}"
                    row_cells[3].text = alignment['alignment_type']
                    row_cells[4].text = str(alignment.get('overlap_details', {}))
            
            # Add footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Generated on: {chart_data['generated_at']}")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            file_path = f"/tmp/chart_{chart_id}.docx"
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error creating DOCX chart: {e}")
            raise
    
    async def create_pdf_chart(self, chart_id: str, chart_data: Dict[str, Any]) -> str:
        """Create a PDF claim chart."""
        try:
            file_path = f"/tmp/chart_{chart_id}.pdf"
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center
            )
            
            # Build content
            story = []
            
            # Add title
            story.append(Paragraph("Claim Chart", title_style))
            story.append(Spacer(1, 20))
            
            # Add patent information
            patent = chart_data['patent']
            claim = chart_data['claim']
            
            story.append(Paragraph("Patent Information", styles['Heading2']))
            patent_text = f"""
            Title: {patent['title']}<br/>
            Publication Number: {patent['pub_number']}<br/>
            Priority Date: {patent['prio_date']}<br/>
            Assignee(s): {', '.join(patent['assignees'])}<br/>
            Inventors: {', '.join(patent['inventors'])}
            """
            story.append(Paragraph(patent_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Add claim text
            story.append(Paragraph(f"Claim {claim['claim_number']}", styles['Heading2']))
            story.append(Paragraph(claim['text'], styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Add novelty information if available
            if chart_data.get('novelty'):
                novelty = chart_data['novelty']
                story.append(Paragraph("Novelty Analysis", styles['Heading2']))
                
                novelty_text = f"""
                Novelty Score: {novelty['novelty_score']:.2f}<br/>
                Obviousness Score: {novelty['obviousness_score']:.2f}<br/>
                Confidence Band: {novelty['confidence_band']}
                """
                story.append(Paragraph(novelty_text, styles['Normal']))
                story.append(Spacer(1, 20))
                
                # Add clause-level details
                if novelty.get('clause_details'):
                    story.append(Paragraph("Clause-Level Analysis", styles['Heading3']))
                    for clause in novelty['clause_details']:
                        clause_text = f"""
                        Clause {clause['clause_index']}: {clause['clause_text']}<br/>
                        Novelty Score: {clause['novelty_score']:.2f} (Confidence: {clause['confidence']})
                        """
                        story.append(Paragraph(clause_text, styles['Normal']))
                        story.append(Spacer(1, 10))
            
            # Add alignments if available
            if chart_data.get('alignments'):
                story.append(Paragraph("Reference Alignments", styles['Heading2']))
                
                # Create alignment table
                table_data = [['Reference Patent', 'Reference Clause', 'Similarity', 'Type', 'Details']]
                
                for alignment in chart_data['alignments']:
                    clause_text = alignment['reference_clause_text'][:50] + '...' if len(alignment['reference_clause_text']) > 50 else alignment['reference_clause_text']
                    table_data.append([
                        alignment.get('reference_patent_title', 'Unknown'),
                        clause_text,
                        f"{alignment['similarity_score']:.3f}",
                        alignment['alignment_type'],
                        str(alignment.get('overlap_details', {}))[:30] + '...'
                    ])
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Add footer
            story.append(Paragraph(f"Generated on: {chart_data['generated_at']}", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error creating PDF chart: {e}")
            raise
    
    async def create_export_bundle(
        self,
        export_id: str,
        patent_ids: List[str],
        export_type: str,
        include_charts: bool,
        include_alignments: bool,
        include_novelty: bool
    ) -> str:
        """Create an export bundle with multiple patents."""
        try:
            import zipfile
            import tempfile
            
            # Create temporary directory for bundle
            with tempfile.TemporaryDirectory() as temp_dir:
                bundle_path = f"/tmp/export_{export_id}.zip"
                
                with zipfile.ZipFile(bundle_path, 'w') as zip_file:
                    # Add summary document
                    summary_path = await self.create_summary_document(
                        temp_dir, patent_ids, include_charts, include_alignments, include_novelty
                    )
                    zip_file.write(summary_path, "summary.pdf")
                    
                    # Add individual patent charts if requested
                    if include_charts:
                        for patent_id in patent_ids:
                            try:
                                # Get first claim for each patent
                                claims = await self.db.get_patent_claims(patent_id)
                                if claims:
                                    claim_num = claims[0]['claim_number']
                                    
                                    # Generate chart data
                                    chart_data = await self.generate_claim_chart(
                                        patent_id, claim_num, include_alignments, include_novelty
                                    )
                                    
                                    # Create chart document
                                    chart_path = await self.create_docx_chart(
                                        f"{export_id}_{patent_id}", chart_data
                                    )
                                    
                                    # Add to zip
                                    zip_file.write(chart_path, f"charts/patent_{patent_id}_claim_{claim_num}.docx")
                                    
                                    # Clean up
                                    os.remove(chart_path)
                                    
                            except Exception as e:
                                logger.warning(f"Failed to create chart for patent {patent_id}: {e}")
                                continue
                
                return bundle_path
                
        except Exception as e:
            logger.error(f"Error creating export bundle: {e}")
            raise
    
    async def create_summary_document(
        self,
        temp_dir: str,
        patent_ids: List[str],
        include_charts: bool,
        include_alignments: bool,
        include_novelty: bool
    ) -> str:
        """Create a summary document for the export bundle."""
        try:
            file_path = os.path.join(temp_dir, "summary.pdf")
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            story.append(Paragraph("Patent Analysis Summary", styles['Heading1']))
            story.append(Spacer(1, 20))
            
            # Add summary information
            story.append(Paragraph(f"Export Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Paragraph(f"Number of Patents: {len(patent_ids)}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Add patent list
            story.append(Paragraph("Patent List", styles['Heading2']))
            
            for i, patent_id in enumerate(patent_ids, 1):
                try:
                    patent = await self.db.get_patent(patent_id)
                    if patent:
                        patent_text = f"{i}. {patent['title']} ({patent['pub_number']})"
                        story.append(Paragraph(patent_text, styles['Normal']))
                except Exception as e:
                    logger.warning(f"Failed to get patent {patent_id}: {e}")
                    continue
            
            # Build PDF
            doc.build(story)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error creating summary document: {e}")
            raise
    
    async def upload_chart_to_storage(self, file_path: str, chart_id: str, chart_type: str) -> str:
        """Upload chart to storage and return URL."""
        try:
            with open(file_path, 'rb') as f:
                file_data = f.read()
            
            # Upload to S3/MinIO
            s3_key = f"charts/{chart_id}.{chart_type}"
            await self.storage.upload_file(s3_key, file_data, f"application/{chart_type}")
            
            # Generate signed URL
            url = await self.storage.get_signed_url(s3_key, expires_in=3600)
            
            return url
            
        except Exception as e:
            logger.error(f"Error uploading chart to storage: {e}")
            raise
    
    async def upload_export_to_storage(self, file_path: str, export_id: str, export_type: str) -> str:
        """Upload export bundle to storage and return URL."""
        try:
            with open(file_path, 'rb') as f:
                file_data = f.read()
            
            # Upload to S3/MinIO
            s3_key = f"exports/{export_id}.{export_type}"
            await self.storage.upload_file(s3_key, file_data, f"application/{export_type}")
            
            # Generate signed URL
            url = await self.storage.get_signed_url(s3_key, expires_in=3600)
            
            return url
            
        except Exception as e:
            logger.error(f"Error uploading export to storage: {e}")
            raise


async def main():
    """Main entry point for the chart worker."""
    worker = ChartWorker()
    
    try:
        await worker.start()
        
        # Keep the worker running
        while True:
            await asyncio.sleep(1)
            
    except KeyboardInterrupt:
        logger.info("Shutting down chart worker...")
    finally:
        await worker.stop()


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    asyncio.run(main())
